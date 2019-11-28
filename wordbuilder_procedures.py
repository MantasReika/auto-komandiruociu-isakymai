import os
import docx
import json
import openpyxl

from docx import Document


def getCountry (cell_country):
    return cell_country.split(" ")[0]
    

def getGalininkaVarda(full_name, linksniai_dic):
    names_split = full_name.strip().split(" ")
    if len(names_split) != 2:
        return False

    
    forename = names_split[1]
    surname = names_split[0]
    try:
        if forename[-1] == "s":
            forename = forename[:-2] + linksniai_dic[forename[-2:]]
        else: 
            forename = forename[:-1] + linksniai_dic[forename[-1:]]

        if surname[-1] == "s":
            surname = surname[:-2] + linksniai_dic[surname[-2:]]
        else: 
            surname =  surname[:-1] + linksniai_dic[surname[-1:]]
    except:
        return False
    return forename + " " + surname


def getXlsxData (file_name, linksniai_dic, INVALID_GA, INVALID_SPEC, COUNTRIES):
    records = []
    unfit_records = []
    START_ROW = 8 #TODO: input from cmd
    FINISH_ROW = START_ROW + int(input("--> Įveskite sveikąjį skaičių: ")) #TODO: automaticly stop after finding blank row
#    INVALID_GA = ["ADM","GTS"]
#    INVALID_SPEC = ["buh.","pavad.","sauga","techn.", "vad.", "vadov.", "vadyb.", "pat"]
#    COUNTRIES = ["Švedija", "Prancūzija", "Belgija", "Šveicarija", "Vokietija", "Latvija", "Ryga", "Lenkija", "Estija", "Nyderlandai", "Olandija", "Suomija"]
    wb = openpyxl.load_workbook(file_name)
    sheet = wb[wb.sheetnames[0]]
    
    for row in range(START_ROW, FINISH_ROW+1):
        this = []
        found_invalid = False
        
        nr = sheet[("A%s" % row)].value
        date_created = sheet[("B%s" % row)].value
        surname_name =  sheet[("C%s" % row)].value
        place = sheet[("D%s" % row)].value
        specialybe = sheet[("E%s" % row)].value
        country = sheet[("F%s" % row)].value
        date_from = sheet[("H%s" % row)].value
        date_to = sheet[("I%s" % row)].value

        if nr != None:

            if surname_name != None:
                surname_name = surname_name.strip()

            if type(nr) == str:
                nr = nr.strip()
            else:
                found_invalid = True
                
            if type(place) == str:
                place = place.strip()
            else:
                found_invalid = True

            if type(specialybe) == str:
                specialybe = specialybe.strip()
            else:
                found_invalid = True

            if type(country) == str:
                country = country.strip()
            else:
                found_invalid = True            
                
            galininkas_full_name = getGalininkaVarda(surname_name, linksniai_dic)
            if len(str(date_from)) > 10:
                date_from = str(date_from)[:10]

            if len(str(date_to)) > 10:
                date_to = str(date_to)[:10]

            if len(str(date_created)) > 10:
                date_created = str(date_created)[:10]


            if galininkas_full_name == False:
                #found_invalid = True
                galininkas_full_name = surname_name

            if country.strip().split(" ")[0] not in COUNTRIES:
                found_invalid = True
      
            this = [nr,date_created, surname_name, place, specialybe, country, date_from, date_to, galininkas_full_name]

            # probably dates are fcing this up. debug if all the field got correct type values
            if not found_invalid:
                for i in this:
                    #if i == "":
                    #    found_invalid = True
                    if i == None:
                        found_invalid = True
                if not found_invalid:
                    if specialybe in INVALID_SPEC:
                        found_invalid = True
                    elif place in INVALID_GA:
                        found_invalid = True
            
            if found_invalid:
                unfit_records.append(this)
            else:
                records.append(this)

            

    print('\n>>> Colected records from xlsx...\n')

    return records, unfit_records

def getDienpinigiai(dienpinigiai_dic, country, value):
    try:    
        if value == 'max':
            return dienpinigiai_dic[country][1]
        else:
            return dienpinigiai_dic[country][0]
    except KeyError:
        return False


class Secondment:
    record = []
    err = False
    
    GALININKO_LINKSNIAI = {}
    MONTHS = {}
    GA1 = {}
    GA2 = {}
    PROFESIJA = {}
    DIENPINIGIAI = {}
    KEYS = {}

    def __init__(self, rec):
        self.record = rec

    def setConstants(self, GALININKO_LINKSNIAI, MONTHS, GA1, GA2, PROFESIJA, DIENPINIGIAI, KEYS):
        self.GALININKO_LINKSNIAI = GALININKO_LINKSNIAI
        self.MONTHS = MONTHS
        self.GA1 = GA1
        self.GA2 = GA2
        self.PROFESIJA = PROFESIJA
        self.DIENPINIGIAI = DIENPINIGIAI
        self.KEYS = KEYS
        
    def translateKeyToValue(self, key):
        if key == '#DOKYEAR#':
            return self.getDocYear()
        elif key == '#DOKMONTH#':
            return self.getDocMonth()
        elif key == '#DOKDAY#':
            return self.getDocDay()
        elif key == '#DOKNUM#':
            return self.getDocNum()
        elif key == '#GA1#':
            return self.getGa1()
        elif key == '#GA2#':
            return self.getGa2()
        elif key == '#CNTRY1#':
            return self.getCountry1()
        elif key == '#CNTRY2#':
            return self.getCountry2()
        elif key == '#KOMANDYEAR#':
            return self.getSecondmentYear()
        elif key == '#KOMANDMONTHDAY#':
            return self.getSecondmentDates()
        elif key == '#PERSONNAME#':
            return self.getPersonName()
        elif key == '#PERSONNAMECASED#':
            return self.getPersonNameCased()
        elif key == '#CNTRY3#':
            return self.geCountry3()
        elif key == '#PROFESSION#':
            return self.getProfesion()
        elif key == '#MONEY1#':
            return self.getAllowence1()
        elif key == '#MONEY2#':
            return self.getAllowence2()
        else:
            raise ValueError("Unexpected Key: ", key)
            #return None
        
    def getDocYear(self):
        return self.record[1].strip()[:4]
    def getDocMonth(self):
        return self.MONTHS[self.record[1].strip()[5:7]]
    def getDocDay(self):
        return self.record[1].strip()[-2:]
    def getDocNum(self):
        return self.record[0].strip()
    def getGa(self):
        return self.record[3].strip()
    def getGa1(self):
        return self.GA1[self.getGa()]
    def getGa2(self):
        return self.GA2[self.getGa()]
    def getCountry(self):
        return self.record[5].strip().split(" ")[0]
    def getCountry1(self):
        return self.getCountry()[:-1] + "os"
    def getCountry2(self):
        return self.getCountry()[:-1] + "ą"
    def getSecondmentYear(self):
        return self.record[6].strip()[:4]
    def getSecondmentMonth1(self):
        return self.MONTHS[self.record[6].strip()[5:7]]
    def getSecondmentMonth2(self):
        return self.MONTHS[self.record[7].strip()[5:7]]
    def getSecondmentDay1(self):
        return self.record[6].strip()[-2:]
    def getSecondmentDay2(self):
        return self.record[7].strip()[-2:]
    def getPersonName(self):
        return self.record[2].strip()
    def getPersonNameCased(self):
        return self.record[8].strip().upper()
    def geCountry3(self):
        return self.getCountry()[:-1] + 'oje'
    def getProfesion(self):
        return self.PROFESIJA[self.record[4].strip()]
    def getAllowence1(self):
        val = getDienpinigiai(self.DIENPINIGIAI, self.getCountry(), "min")
        if val == False:
            self.err = "FAIL"
        else:
            return val
    def getAllowence2(self):
        val = getDienpinigiai(self.DIENPINIGIAI, self.getCountry(), "max")
        if val == False:
            self.err = "FAIL"
        else:
            return val
    def getSecondmentDates(self):
        if (self.getSecondmentMonth1() == self.getSecondmentMonth2()) and (self.getSecondmentDay1() == self.getSecondmentDay2()):
            return "%s %s dieną" % (self.getSecondmentMonth1(), self.getSecondmentDay1())
        elif (self.getSecondmentMonth1() == self.getSecondmentMonth2()) and (self.getSecondmentDay1() != self.getSecondmentDay2()):
            return "%s %s-%s dienomis imt." % (self.getSecondmentMonth1(), self.getSecondmentDay1(), self.getSecondmentDay2())
        else:
            return "%s %s - %s %s dienomis imt." % (self.getSecondmentMonth1(), self.getSecondmentDay1(), self.getSecondmentMonth2(), self.getSecondmentDay2())

def indexParagrahps(c):
    doc = Document(c.TEMPLATE_NAME)
    paragraphIdx = {}
    for pNum in range(len(doc.paragraphs)):
        for key in c.KEYS:
            if key in doc.paragraphs[pNum].text:
                paragraphIdx[key] = pNum
    return paragraphIdx

def indexRuns(c, paragraphIdx):
    doc = Document(c.TEMPLATE_NAME)
    runIdx = {}
    for key in c.KEYS:
        for rNum in range(len(doc.paragraphs[paragraphIdx[key]].runs)):
            if key in doc.paragraphs[paragraphIdx[key]].runs[rNum].text:
                runIdx[key] = rNum
    return runIdx

def createSecondmentDocument (record, doc, paragraphIdx, runIdx, c, goodRecord):
    if goodRecord:
        docDir = c.DOCUMENTS_DIR
    else:
        docDir = c.FAILED_DOCUMENTS_DIR
    
    #doc = Document(c.TEMPLATE_NAME)
    s = Secondment(record)
    s.setConstants(c.GALININKO_LINKSNIAI, c.MONTHS, c.GA1, c.GA2, c.PROFESIJA, c.DIENPINIGIAI, c.KEYS)

    try:
        for key in c.KEYS:
            placeHolderRun = doc.paragraphs[paragraphIdx[key]].runs[runIdx[key]].text
    #        print("%s %s" % (key, s.translateKeyToValue(key)))
            doc.paragraphs[paragraphIdx[key]].runs[runIdx[key]].text = placeHolderRun.replace(key, s.translateKeyToValue(key))
    except Exception as e:
        print('Something unrecognised for: %s\nError message: %s\n' % (record[2], e))
        return
    currentDir = os.getcwd()
    
    if currentDir[-len(s.getCountry()):] != s.getCountry().upper():       
        try:
            os.chdir(c.FILE_DIR + '/' + docDir)
            os.chdir(s.getCountry().upper())
        except FileNotFoundError:
            os.chdir(c.FILE_DIR + '/' + docDir)
            os.mkdir(s.getCountry().upper())
            os.chdir(s.getCountry().upper())
    fileNr = ""
    fileExists = os.path.isfile(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx') 
    if not fileExists:
        doc.save(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx') 
    else:
        fileNr = 0
        while fileExists:
            fileNr += 1 
            fileExists = os.path.isfile(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx') 
            if not fileExists:
                doc.save(s.getPersonName().split(' ')[0] + " " + s.getGa() + " " + "%s" % fileNr + '.docx')

def loadJsonData():
    with open('GALININKO_LINKSNIAI.json', 'r', encoding='utf8') as fp:
        GALININKO_LINKSNIAI = json.load(fp)
        
    with open('MONTHS.json', 'r', encoding='utf8') as fp:
        MONTHS = json.load(fp)

    with open('GA1.json', 'r', encoding='utf8') as fp:
        GA1 = json.load(fp)

    with open('GA2.json', 'r', encoding='utf8') as fp:
        GA2 = json.load(fp)

    with open('PROFESIJA.json', 'r', encoding='utf8') as fp:
        PROFESIJA = json.load(fp)

    with open('DIENPINIGIAI.json', 'r', encoding='utf8') as fp:
        DIENPINIGIAI = json.load(fp)
        
    with open('DocumentPlaceHolders.json', 'r', encoding='utf8') as fp:
        KEYS = json.load(fp)

    return GALININKO_LINKSNIAI, MONTHS, GA1, GA2, PROFESIJA, DIENPINIGIAI, KEYS

class Constants:
    TEMPLATE_NAME = ""
    INVALID_GA = []
    INVALID_SPEC = []
    COUNTRIES = []
    TIKSLAS = ""
    DOCUMENTS_DIR = ""
    FAILED_DOCUMENTS_DIR = ""
    FILE_DIR = ""
    GALININKO_LINKSNIAI = {}
    MONTHS = {}
    GA1 = {}
    GA2 = {}
    PROFESIJA = {}
    DIENPINIGIAI = {}
    KEYS = {}

    def __init__(self):
        self.TEMPLATE_NAME = "template.docx"
        self.INVALID_GA = ["ADM","GTS"]
        self.INVALID_SPEC = ["pavad.","sauga","techn.", "vadov.", "pat"]
        self.COUNTRIES = ["Švedija", "Prancūzija", "Belgija", "Šveicarija", "Vokietija", "Latvija", "Ryga", "Lenkija", "Estija", "Nyderlandai", "Olandija", "Suomija", "Danija"]
        self.TIKSLAS = "Atlikti paskirtus darbus objekte %s pagal sutartį."
        self.DOCUMENTS_DIR = "Komandiruočių Įsakymai"
        self.FAILED_DOCUMENTS_DIR = "Negarantuoti Komandiruočių Įsakymai"
        self.FILE_DIR = os.getcwd()
        self.GALININKO_LINKSNIAI, self.MONTHS, self.GA1, self.GA2, self.PROFESIJA, self.DIENPINIGIAI, self.KEYS = loadJsonData()
